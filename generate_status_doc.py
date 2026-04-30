from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_status_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('IDP Pipeline: Project Status Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction
    doc.add_paragraph("This document provides a detailed summary of the progress made on the Pharmaceutical Blister Pack Defect Detection system and outlines the roadmap for the final integration phases.")
    
    # Section 1: Accomplished Tasks
    doc.add_heading('1. Accomplished Tasks', level=1)
    
    # YOLO
    p1 = doc.add_paragraph()
    p1.add_run('YOLOv8 Frontside Detection: ').bold = True
    p1.add_run('Successfully trained a YOLOv8-nano model to identify Good Packs, Defects, and Missing Packs. Validation metrics and confusion matrices have been generated.')
    
    # OCR
    p2 = doc.add_paragraph()
    p2.add_run('OCR Validation Pipeline: ').bold = True
    p2.add_run('Integrated PaddleOCR (v2.7.0.3) with a specialized preprocessing suite (CLAHE, Bilateral Filtering, 3x Upscaling, and Image Inversion) to handle challenging metallic foil surfaces. Successfully generated structured JSON results for 101 validation images.')
    
    # Section 2: Phase 1 — PatchCore Anomaly Detection
    doc.add_heading('2. Phase 1: PatchCore (In Progress)', level=1)
    doc.add_paragraph("Goal: Implement unsupervised anomaly detection for the blister pack backside using the Anomalib library.")
    
    tasks_pc = [
        "Data Preparation: Isolated 474 'normal' training samples and balanced the test set (50 good / 49 defect).",
        "Training Script: Developed 'train_patchcore.py' using a ResNet18 backbone optimized for GPU training.",
        "Status: Environment is configured; training execution is the immediate next step."
    ]
    for task in tasks_pc:
        doc.add_paragraph(task, style='List Bullet')
        
    # Section 3: Remaining Roadmap
    doc.add_heading('3. Remaining Roadmap', level=1)
    
    # Phase 2
    doc.add_heading('Phase 2: Federated Learning Simulation', level=2)
    doc.add_paragraph("Simulate a collaborative training environment across 3 pharmaceutical plant sites (Site A, B, and C).")
    tasks_fl = [
        "Data Splitting: Divide 1014 training images into 3 balanced sites (~158 good + 180 defect each).",
        "Orchestration: Set up a Flower (flwr) server and 3 client processes.",
        "Training: Run 5 rounds of FedAvg training to demonstrate model convergence without raw data sharing."
    ]
    for task in tasks_fl:
        doc.add_paragraph(task, style='List Bullet')
        
    # Phase 3
    doc.add_heading('Phase 3: Unified Streamlit Dashboard', level=2)
    doc.add_paragraph("Create a comprehensive 4-page monitoring application:")
    tasks_db = [
        "Page 1: Frontside YOLO Detection results and real-time inference.",
        "Page 2: Backside OCR Validation (3-state logic: PASS, UNREADABLE, REJECT).",
        "Page 3: PatchCore Anomaly Heatmaps and score distributions.",
        "Page 4: Federated Learning convergence and privacy compliance metrics."
    ]
    for task in tasks_db:
        doc.add_paragraph(task, style='List Bullet')

    # Save
    filename = "IDP_Project_Status_Report.docx"
    doc.save(filename)
    print(f"Status report generated: {filename}")

if __name__ == "__main__":
    create_status_report()
