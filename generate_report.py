from docx import Document
from docx.shared import Inches

def create_report():
    # Initialize a blank Word Document
    doc = Document()
    doc.add_heading('Pharmaceutical Defect Detection — Model Accuracy & Confusion Matrices', 0)

    # Intro
    doc.add_paragraph(
        "Attached below are the normalized confusion matrices for both Camera 1 and Camera 2 models, "
        "generated natively by the YOLOv8 validation process. The dark blue diagonal "
        "axis from top-left to bottom-right represents the True Positive Rate."
    )

    # ------------------ Camera 1 ------------------
    doc.add_heading('Camera 1: Frontside Model Accuracy', level=1)
    doc.add_paragraph(
        "True Positive Rates (Diagonal): Most classes achieve ~90% to 100% exact matches "
        "with the ground truth labels. The background class causes very few false positives."
    )
    try:
        doc.add_picture('C:\\temp\\pharma_defect\\yolov8n_blister\\confusion_matrix_normalized.png', width=Inches(6.0))
        doc.add_paragraph("Figure 1: Normalized Confusion Matrix - Camera 1")
    except Exception as e:
        doc.add_paragraph(f"[Failed to load Camera 1 Image: {e}]")

    # ------------------ Camera 2 ------------------
    doc.add_heading('Camera 2: Backside Model Accuracy', level=1)
    
    # Bullet points for performance
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run('Defective packs classified correctly 95% of the time.')
    
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run('Good packs classified correctly 100% of the time.')

    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run('The no_pack class had an extremely low sample rate in the test pipeline, giving it lower confidence natively, but standard detection safely separates good_pack and defect.')

    try:
        doc.add_picture('C:\\temp\\pharma_defect\\yolov8n_backside2\\confusion_matrix_normalized.png', width=Inches(6.0))
        doc.add_paragraph("Figure 2: Normalized Confusion Matrix - Camera 2")
    except Exception as e:
        doc.add_paragraph(f"[Failed to load Camera 2 Image: {e}]")

    # Save Document
    doc_path = 'Model_Performance_Report.docx'
    doc.save(doc_path)
    print(f"Report saved as: {doc_path}")

if __name__ == "__main__":
    create_report()
